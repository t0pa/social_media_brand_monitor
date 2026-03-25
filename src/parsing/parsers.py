import os
import json
import csv
import xml.etree.ElementTree as ET
import sys
import pdfplumber
import re
from docx import Document
from openpyxl import load_workbook
import chardet

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'storage')))
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))
from src.storage.mongo import save_to_mongo
from src.utils.logger import logging

# --- Text Normalization ---
def normalize_text(text):
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()

# --- File Reading with Encoding Detection ---
def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()
    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    logging.info(f"Detected encoding for {os.path.basename(file_path)}: {encoding} (confidence: {result.get('confidence')})")
    return raw.decode(encoding, errors="replace")

# --- Existing Parsers (JSON, CSV, XML) ---
def parse_json_files(directory):
    logging.info(f"Parsing JSON files from {directory}")
    for filename in os.listdir(directory):
        if filename.endswith(".json"):
            filepath = os.path.join(directory, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    data = json.load(f)
                metadata = {"source": filename, "document_type": "json"}
                if isinstance(data, dict) and "articles" in data:
                    save_to_mongo(data["articles"], metadata=metadata)
                    logging.info(f"Saved {len(data['articles'])} articles from {filename}")
                else:
                    save_to_mongo(data, metadata=metadata)
                    logging.info(f"Saved data from {filename}")
            except Exception as e:
                logging.error(f"Failed to parse JSON file {filename}: {e}")

def parse_csv_file(filepath):
    logging.info(f"Parsing CSV file: {os.path.basename(filepath)}")
    try:
        with open(filepath, newline='', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            rows = list(reader)
        metadata = {"source": os.path.basename(filepath), "document_type": "csv"}
        save_to_mongo(rows, metadata=metadata)
        logging.info(f"Saved {len(rows)} rows from {os.path.basename(filepath)}")
    except Exception as e:
        logging.error(f"Failed to parse CSV file {os.path.basename(filepath)}: {e}")

def parse_xml_file(filepath):
    logging.info(f"Parsing XML file: {os.path.basename(filepath)}")
    try:
        tree = ET.parse(filepath)
        root = tree.getroot()
        data = [{child.tag: child.text for child in elem} for elem in root]
        metadata = {"source": os.path.basename(filepath), "document_type": "xml"}
        save_to_mongo(data, metadata=metadata)
        logging.info(f"Saved {len(data)} elements from {os.path.basename(filepath)}")
    except Exception as e:
        logging.error(f"Failed to parse XML file {os.path.basename(filepath)}: {e}")

# --- PDF Parsers ---
def extract_text_from_pdf(pdf_path):
    """Extracts all text from a standard PDF file."""
    with pdfplumber.open(pdf_path) as pdf:
        return "\n\n".join(normalize_text(page.extract_text()) for page in pdf.pages if page.extract_text())

def extract_text_from_two_column_pdf(pdf_path, gap=10):
    """Extracts text from a two-column PDF layout."""
    pages_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            mid_x = page.width / 2
            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))
            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")
            combined = "\n\n".join(part for part in [left_text, right_text] if part)
            if combined:
                pages_text.append(combined)
    return "\n\n".join(pages_text)

def extract_tables_from_pdf(pdf_path):
    """Extracts all tables from a PDF file."""
    with pdfplumber.open(pdf_path) as pdf:
        return [page.extract_tables() for page in pdf.pages]

def parse_pdf_files(directory):
    logging.info(f"Parsing PDF files from {directory}")
    for filename in os.listdir(directory):
        if filename.endswith(".pdf"):
            filepath = os.path.join(directory, filename)
            logging.info(f"Processing PDF: {filename}")
            try:
                with pdfplumber.open(filepath) as pdf:
                    metadata = {"source": filename, "document_type": "pdf", **pdf.metadata}

                if "columns" in filename.lower():
                    text = extract_text_from_two_column_pdf(filepath)
                else:
                    text = extract_text_from_pdf(filepath)
                
                tables = extract_tables_from_pdf(filepath)
                
                if text:
                    save_to_mongo({"type": "text", "content": text}, metadata=metadata)
                
                for i, page_tables in enumerate(tables):
                    if page_tables:
                        save_to_mongo([{"type": "table", "page": i+1, "content": table} for table in page_tables], metadata=metadata)
            except Exception as e:
                logging.error(f"Failed to process PDF {filename}: {e}")

# --- Word (DOCX) Parser ---
def extract_text_from_word(docx_path):
    """Extracts all text from paragraphs in a docx file."""
    doc = Document(docx_path)
    return "\n\n".join(normalize_text(para.text) for para in doc.paragraphs if para.text)

def extract_tables_from_word(docx_path):
    """Extracts all tables from a docx file."""
    doc = Document(docx_path)
    all_tables = []
    for table in doc.tables:
        # Extracting header and rows
        header = [cell.text for cell in table.rows[0].cells]
        data_rows = []
        for row in table.rows[1:]:
            data_rows.append({header[i]: cell.text for i, cell in enumerate(row.cells)})
        all_tables.append(data_rows)
    return all_tables

def parse_docx_files(directory):
    logging.info(f"Parsing DOCX files from {directory}")
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            filepath = os.path.join(directory, filename)
            logging.info(f"Processing DOCX: {filename}")
            try:
                metadata = {"source": filename, "document_type": "docx"}
                
                text = extract_text_from_word(filepath)
                if text:
                    save_to_mongo({"type": "text", "content": text}, metadata=metadata)
                
                tables = extract_tables_from_word(filepath)
                if tables:
                    save_to_mongo([{"type": "table", "content": table} for table in tables], metadata=metadata)
            except Exception as e:
                logging.error(f"Failed to process DOCX {filename}: {e}")

# --- Excel (XLSX) Parser ---
def extract_data_from_excel_sheet(file_path, sheet_name):
    """Extracts data from a specific sheet of an Excel file as a list of dictionaries."""
    wb = load_workbook(file_path, data_only=True)
    ws = wb[sheet_name]
    
    # Assume first row is header
    header = [cell.value for cell in ws[1]]
    data = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(cell is not None for cell in row):
            data.append(dict(zip(header, row)))
    return data

def parse_excel_files(directory):
    logging.info(f"Parsing Excel files from {directory}")
    for filename in os.listdir(directory):
        if filename.endswith((".xlsx", ".xls")):
            filepath = os.path.join(directory, filename)
            logging.info(f"Processing Excel: {filename}")
            try:
                wb = load_workbook(filepath, data_only=True)
                metadata = {"source": filename, "document_type": "excel"}
                all_sheets_data = {}
                for sheetname in wb.sheetnames:
                    sheet_data = extract_data_from_excel_sheet(filepath, sheetname)
                    if sheet_data:
                        all_sheets_data[sheetname] = sheet_data
                
                if all_sheets_data:
                    save_to_mongo(all_sheets_data, metadata=metadata)
            except Exception as e:
                logging.error(f"Failed to process Excel {filename}: {e}")

# --- Main Execution ---
if __name__ == "__main__":
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    
    # Define directories
    api_dir = os.path.join(project_root, "data", "raw", "api")
    csv_dir = os.path.join(project_root, "data", "raw", "csv")
    xml_dir = os.path.join(project_root, "data", "raw", "xml")
    pdf_dir = os.path.join(project_root, "data", "raw", "pdf")
    docx_dir = os.path.join(project_root, "data", "raw", "docx")
    excel_dir = os.path.join(project_root, "data", "raw", "excel")

    # Run parsers
    if os.path.exists(api_dir): parse_json_files(api_dir)
    if os.path.exists(csv_dir):
        for f in os.listdir(csv_dir):
            if f.endswith(".csv"): parse_csv_file(os.path.join(csv_dir, f))
    if os.path.exists(xml_dir):
        for f in os.listdir(xml_dir):
            if f.endswith(".xml"): parse_xml_file(os.path.join(xml_dir, f))
    if os.path.exists(pdf_dir): parse_pdf_files(pdf_dir)
    if os.path.exists(docx_dir): parse_docx_files(docx_dir)
    if os.path.exists(excel_dir): parse_excel_files(excel_dir)
    
    logging.info("All parsing complete.")
